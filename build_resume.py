"""Generate a polished, ATS-friendly resume for Art Van de Riet."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---- Palette ----
NAVY = RGBColor(0x1F, 0x2A, 0x44)      # deep navy for name / headings
ACCENT = RGBColor(0x2E, 0x5E, 0x8C)    # steel blue accent
DARK = RGBColor(0x22, 0x22, 0x22)      # body text
GRAY = RGBColor(0x55, 0x55, 0x55)      # subtitles / dates

doc = Document()

# ---- Base document styling ----
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10.5)
normal.font.color.rgb = DARK
pf = normal.paragraph_format
pf.space_after = Pt(0)
pf.space_before = Pt(0)
pf.line_spacing = 1.06

# ---- Margins ----
for section in doc.sections:
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)


def set_space(p, before=0, after=0):
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)


def add_bottom_border(paragraph, color="2E5E8C", size="8"):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), color)
    pbdr.append(bottom)
    pPr.append(pbdr)


def section_heading(text):
    p = doc.add_paragraph()
    set_space(p, before=10, after=3)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(11.5)
    run.font.color.rgb = NAVY
    run.font.name = "Calibri"
    # letter spacing
    rPr = run._element.get_or_add_rPr()
    spc = OxmlElement("w:spacing")
    spc.set(qn("w:val"), "30")
    rPr.append(spc)
    add_bottom_border(p)
    return p


def bullet(text, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet")
    set_space(p, before=0, after=2)
    p.paragraph_format.left_indent = Inches(0.22)
    p.paragraph_format.first_line_indent = Inches(-0.16)
    if bold_lead:
        r = p.add_run(bold_lead)
        r.bold = True
        r.font.color.rgb = DARK
    r2 = p.add_run(text)
    r2.font.color.rgb = DARK
    return p


# =====================================================================
# HEADER
# =====================================================================
name_p = doc.add_paragraph()
name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_space(name_p, after=0)
nr = name_p.add_run("ART VAN DE RIET")
nr.bold = True
nr.font.size = Pt(26)
nr.font.color.rgb = NAVY
rPr = nr._element.get_or_add_rPr()
spc = OxmlElement("w:spacing")
spc.set(qn("w:val"), "40")
rPr.append(spc)

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_space(title_p, before=1, after=3)
tr = title_p.add_run("Systems Engineer  |  Automation & Change Management")
tr.font.size = Pt(11.5)
tr.font.color.rgb = ACCENT
tr.bold = True

contact_p = doc.add_paragraph()
contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_space(contact_p, before=0, after=2)
contact_bits = [
    "Imperial, MO",
    "314-574-8894",
    "Arthur.vanderiet@gmail.com",
    "linkedin.com/in/art-van-de-riet-75188155",
]
cr = contact_p.add_run("   \u2022   ".join(contact_bits))
cr.font.size = Pt(9.5)
cr.font.color.rgb = GRAY

# thin rule under header
rule = doc.add_paragraph()
set_space(rule, before=0, after=0)
add_bottom_border(rule, color="1F2A44", size="12")

# =====================================================================
# PROFESSIONAL SUMMARY
# =====================================================================
section_heading("Professional Summary")
summary = doc.add_paragraph()
set_space(summary, before=2, after=2)
summary.add_run(
    "Results-driven systems engineer with 10+ years of experience spanning network "
    "operations, service delivery, and change management. Specializes in building "
    "automation and reporting solutions with Python, the Microsoft Power Platform, and "
    "SQL that eliminate process bottlenecks, prevent service-impacting outages, and give "
    "technical teams faster, data-driven decisions. Recognized for delivering "
    "production-ready tools, high-impact executive reporting, and process improvements "
    "that protect large-scale networks."
)

# =====================================================================
# CORE SKILLS
# =====================================================================
section_heading("Core Skills")

skill_groups = [
    ("Automation & Development", "Python  \u00b7  Power Automate  \u00b7  Power Apps  \u00b7  VBA / Excel Macros  \u00b7  C# / .NET  \u00b7  Azure Functions  \u00b7  Microsoft Graph API  \u00b7  REST APIs  \u00b7  SharePoint  \u00b7  Git  \u00b7  CI/CD & GitLab Runners  \u00b7  pytest"),
    ("Cloud & AWS", "AWS  \u00b7  EC2  \u00b7  S3  \u00b7  VPC  \u00b7  IAM  \u00b7  Lambda  \u00b7  RDS  \u00b7  DynamoDB  \u00b7  Route 53  \u00b7  CloudFront  \u00b7  Auto Scaling  \u00b7  ELB  \u00b7  SNS  \u00b7  SQS  \u00b7  CloudWatch  \u00b7  CloudFormation"),
    ("Data & Reporting", "SQL  \u00b7  T-SQL  \u00b7  pandas  \u00b7  openpyxl  \u00b7  ETL / Data Pipelines  \u00b7  Data Modeling  \u00b7  CMDB  \u00b7  Power BI  \u00b7  Tableau  \u00b7  Microsoft Excel & Access  \u00b7  Executive Reporting"),
    ("Operations & Process", "ITIL Framework  \u00b7  Change Management  \u00b7  Incident Management  \u00b7  Problem Management  \u00b7  Root Cause Analysis  \u00b7  Network Surveillance  \u00b7  Method of Procedure (MOP)  \u00b7  Compliance & Risk"),
    ("Platforms & Tools", "Microsoft 365  \u00b7  Azure  \u00b7  SharePoint Online  \u00b7  Visio  \u00b7  GitLab  \u00b7  BMC Remedy  \u00b7  Jira  \u00b7  Netcool  \u00b7  PowerPoint"),
]
for label, items in skill_groups:
    p = doc.add_paragraph()
    set_space(p, before=1, after=1)
    p.paragraph_format.left_indent = Inches(0.02)
    lr = p.add_run(label + ":  ")
    lr.bold = True
    lr.font.color.rgb = ACCENT
    ir = p.add_run(items)
    ir.font.color.rgb = DARK

# =====================================================================
# PROFESSIONAL EXPERIENCE
# =====================================================================
section_heading("Professional Experience")


def job(title, company, location, dates):
    p = doc.add_paragraph()
    set_space(p, before=7, after=0)
    # tab stop at right margin for the date
    p.paragraph_format.tab_stops.add_tab_stop(Inches(7.1), WD_TAB_ALIGNMENT.RIGHT)
    tr = p.add_run(title)
    tr.bold = True
    tr.font.size = Pt(11)
    tr.font.color.rgb = NAVY
    dr = p.add_run("\t" + dates)
    dr.bold = True
    dr.font.size = Pt(9.5)
    dr.font.color.rgb = GRAY
    # company line
    cp = doc.add_paragraph()
    set_space(cp, before=0, after=2)
    ccr = cp.add_run(company + "  |  " + location)
    ccr.italic = True
    ccr.font.size = Pt(9.5)
    ccr.font.color.rgb = ACCENT


job("Systems Engineer II \u2014 Automation & Network Intelligence", "Spectrum (Charter Communications)", "Town and Country, MO", "Apr 2026 \u2013 Present")
bullet("Architected and shipped the CM Hub Intel Tool, an end-to-end network intelligence platform that consolidates 3,916 hub records from 6 source systems into a single searchable CMDB used daily by NOC and Change Management teams.")
bullet("Wrote T-SQL against the Remedy (NOC_APPS) database to pull hub, node, and device (CMTS/OLT/RDM) data, then built a Python (pandas / openpyxl) pipeline that merges optical transport, single-threaded, and backbone sources into one dataset.")
bullet("Developed a two-pass matching engine (exact + vendor-locked fuzzy matching) that links 100+ Visio transport maps to hubs at a 96% match rate, and parsed Visio .vsdx XML to flag 167 backbone (BBoC) hubs and 108 single-threaded sites.")
bullet("Built a C# / .NET Microsoft Graph API updater with certificate-based auth, batched parallel writes, and exponential-backoff retry that replaced a 2-hour Power Automate flow with a 5\u201310 minute upload \u2014 cutting the full update cycle from ~2.5 hours to ~20 minutes.")
bullet("Delivered the tool through a Power Apps front end on a SharePoint list, added a serverless C# Azure Function for automated Visio splitting, and set up GitLab CI/CD with property-based (Hypothesis) tests to protect data-integrity invariants.")

job("Change Analyst II \u2014 Change Management & Automation", "Spectrum (Charter Communications)", "Town and Country, MO", "2018 \u2013 Apr 2026")
bullet("Review 1,000+ planned maintenance changes annually to validate scope, verify device relationships and impacts, and prevent self-induced subscriber outages.")
bullet("Approved 25% of all Inside Plant changes in 2022 \u2014 the company's largest organization \u2014 with 0 administrative errors and 0 break/fix incidents on approved tickets.")
bullet("Received the \u201cAbove and Beyond\u201d award (Q3 2022) for rebuilding and automating department-wide daily, monthly, and yearly reporting.")
bullet("Build and present monthly executive reporting and year-over-year trend analysis that guides leadership planning and improves department efficiency.")
bullet("Designed a standardized messaging application (VBA/Excel) that streamlined maintenance communications and cut email preparation time.")
bullet("Authored method-of-procedure templates and process documentation across 7 verticals, improving CRQ accuracy and reducing escalations with Inside Plant leadership.")

job("Associate Network Operations", "Spectrum (Charter Communications)", "Town and Country, MO", "2016 \u2013 2018")
bullet("Monitored network events with Netcool and BMC Remedy to triage, escalate, and resolve critical outages across VoIP, HSD, and Video services.")
bullet("Led incident management conference calls, coordinated multi-team responses, and delegated efforts to reduce average resolution time.")
bullet("Proactively investigated alarms to minimize customer impact and helped launch the new Remedy ticketing platform, training peers on adoption.")

job("Service Delivery Coordinator", "Spectrum (Charter Communications)", "St. Louis, MO", "2015 \u2013 2016")
bullet("Resolved provisioning and billing errors in CSG systems, specializing in E911 and number porting accuracy.")
bullet("Identified and corrected recurring order-entry issues, preventing thousands in potential billing losses.")

job("Telephone Service Advisor I / II / Sr. Advisor", "Spectrum (Charter Communications)", "St. Louis, MO", "2013 \u2013 2015")
bullet("Performed complex troubleshooting across phone, internet, and cable services while consistently meeting leadership performance metrics.")
bullet("Served as floor support for the supervisor and mentored new hires, fostering a collaborative, high-performing team environment.")

# =====================================================================
# KEY PROJECTS
# =====================================================================
section_heading("Key Projects")


def project(title, tech, lines):
    p = doc.add_paragraph()
    set_space(p, before=6, after=1)
    tr = p.add_run(title)
    tr.bold = True
    tr.font.size = Pt(10.5)
    tr.font.color.rgb = NAVY
    if tech:
        tt = p.add_run("   \u2014  " + tech)
        tt.italic = True
        tt.font.size = Pt(9)
        tt.font.color.rgb = GRAY
    for line in lines:
        bullet(line)


project(
    "CM Hub Intel Tool", "Python, pandas, Power Apps, Power Automate, SharePoint, C# / Azure Functions, GitLab CI/CD",
    [
        "Built an end-to-end network intelligence platform that consolidates 3,916 hub records from 6 source systems, giving NOC and Change Management teams instant search access to transport maps, backbone dependencies, and risk indicators.",
        "Engineered a Python pipeline that parses 100+ Visio network maps and links them to hubs at a 96% match rate, replacing manual searches through a 100+ page PDF and cutting hub lookups from minutes to seconds.",
        "Flagged 167 backbone (BBoC) hubs and 108 single-threaded sites with automated warning banners, surfacing outage blast radius at the point of CRQ review to prevent service-impacting changes.",
        "Delivered the tool through a Power Apps front end backed by a SharePoint list, with a C# Azure Function automating Visio map splitting and a GitLab CI/CD pipeline.",
    ],
)
project(
    "STL Hub 72-Hour Buffer Analyzer", "Python, pandas, Excel, OneDrive",
    [
        "Built a Python tool that monitors 20 St. Louis hub locations to detect and prevent maintenance conflicts using a 72-hour buffer rule.",
        "Auto-detects Excel files from OneDrive and generates color-coded reports flagging buffer violations, active/expired buffers, and compliant tasks.",
        "Groups conflicting work into Multi-Task Resolution (MTR) groups and produces automated reschedule recommendations, with batch and drag-and-drop launchers.",
    ],
)
project(
    "Conflict Assessment & Notification System", "Power Automate, Power Apps, SharePoint, Excel",
    [
        "Developed an automated workflow that detects CRQs with overlapping transport paths and stores grouped conflict records with dynamic Conflict IDs.",
        "Delivered a Power Apps dashboard letting stakeholders view grouped records, mark resolutions, and track live status, with automated email alerts to coordinators.",
    ],
)
project(
    "Transport Lookup Application", "SharePoint, Power Automate, Access",
    [
        "Designed a SharePoint-integrated lookup system that identifies CRQs referencing the same transport map, enabling technicians to self-serve conflict insights and cutting turnaround time.",
    ],
)
project(
    "Centralized Communication Application", "VBA, Excel",
    [
        "Created a centralized tool that standardized maintenance communication verbiage, improving consistency and reducing time spent on outbound notifications.",
    ],
)

# =====================================================================
# EDUCATION & CERTIFICATIONS
# =====================================================================
section_heading("Education & Certifications")

edu = doc.add_paragraph()
set_space(edu, before=3, after=1)
er = edu.add_run("Associate of Applied Science, Information Technology & Computer Network Systems")
er.bold = True
er.font.color.rgb = NAVY
edu2 = doc.add_paragraph()
set_space(edu2, before=0, after=3)
er2 = edu2.add_run("ITT Technical Institute, Arnold, MO  \u00b7  GPA 3.2  \u00b7  2011\u20132013")
er2.italic = True
er2.font.size = Pt(9.5)
er2.font.color.rgb = GRAY

bullet("Leadership Development: Charter Communications Technical Leadership Development Program (TLDP) \u2014 selected participant, 2026 cohort (8-month cross-functional program: mentoring, DDI leadership modules, and an executive presentation to senior leadership).", bold_lead="")
bullet("SCTE Certifications: DOCSIS Engineering Professional, Broadband Telecommunications Specialist, Understanding Cable Technology.", bold_lead="")
bullet("AWS Training: AWS Cloud Quest \u2014 Cloud Practitioner (completed); Solutions Architect (in progress).", bold_lead="")
bullet("Professional Development (Udemy): Python (Beginner to Advanced), SQL, Microsoft Access (Beginner to Advanced), Tableau Data Analytics, Analytics Engineering Bootcamp.", bold_lead="")

# ---- Save ----
out = r"c:\Users\avanderiet\Resume\Art_Van_de_Riet_Resume_2026.docx"
doc.save(out)
print("Saved:", out)
